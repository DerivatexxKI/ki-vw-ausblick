import streamlit as st
import os
from openai import OpenAI
from docx import Document
from io import BytesIO
import pandas as pd
import pdfplumber

# OpenAI API-Key prüfen
api_key = os.getenv("OPENAI_API_KEY")

# Debug-Anzeige
st.write("✅ API-Key geladen:", api_key[:8] + "..." if api_key else "❌ NICHT gesetzt")

if not api_key:
    st.error("❌ Kein API-Key gefunden. Bitte in GitHub unter Settings > Secrets > Actions den Key als OPENAI_API_KEY eintragen.")
    st.stop()

# OpenAI initialisieren
client = OpenAI(api_key=api_key)

# Streamlit UI konfigurieren
st.set_page_config(page_title="📊 Volkswirtschaftlicher Ausblick Generator", page_icon="📈")
st.title("📈 KI-gestützter Volkswirtschaftlicher Ausblick")

# Dateiuploads
uploaded_excels = st.file_uploader("📥 Excel-Dateien mit Zinsdaten hochladen", type=["xlsx", "xls", "csv"], accept_multiple_files=True)
uploaded_pdfs = st.file_uploader("📥 PDF-Dateien mit Research- oder Marktberichten hochladen", type="pdf", accept_multiple_files=True)

# Funktion zum PDF-Text extrahieren
def extract_text_from_pdfs(pdfs):
    text = ""
    for pdf in pdfs:
        with pdfplumber.open(pdf) as pdf_reader:
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
    return text

# Excel-Zusammenfassung
def summarize_excels(files):
    summaries = []
    for f in files:
        try:
            df = pd.read_excel(f)
        except:
            df = pd.read_csv(f)
        summaries.append(f"Datei: {f.name}\nSpalten: {', '.join(df.columns)}\nVorschau:\n{df.head(3).to_string(index=False)}")
    return "\n\n".join(summaries)

# Knopf zur Erstellung des Berichts
if st.button("🧠 Bericht erstellen"):
    with st.spinner("Analysiere Inhalte und erstelle Bericht..."):

        pdf_text = extract_text_from_pdfs(uploaded_pdfs) if uploaded_pdfs else ""
        excel_summary = summarize_excels(uploaded_excels) if uploaded_excels else ""

        prompt = f"""
        Du bist ein Experte für Volkswirtschaft, Zentralbankpolitik und Zinsprognosen.
        Erstelle einen professionellen, strukturierten volkswirtschaftlichen Ausblick für eine deutsche Regionalbank.
        Verwende diese Inhalte aus PDFs:
        {pdf_text[:6000]}
        
        Und diese Excel-Daten:
        {excel_summary[:3000]}

        Struktur des Ausblicks:
        1. Aktuelle wirtschaftliche Lage
        2. Zinsumfeld (EZB, FED, Markt)
        3. Inflationsausblick
        4. Risiken & Unsicherheiten
        5. Mittelfristiger Ausblick und Zinsprojektion
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )

        report_text = response.choices[0].message.content

        # Word-Dokument generieren
        doc = Document()
        doc.add_heading("Volkswirtschaftlicher Ausblick", 0)
        for line in report_text.split("\n"):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Bericht erstellt!")
        st.download_button(
            "📄 Bericht als Word-Datei herunterladen",
            buffer,
            file_name="Volkswirtschaftlicher_Ausblick.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
